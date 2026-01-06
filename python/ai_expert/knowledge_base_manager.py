
# -*- coding: utf-8 -*-
import os
import json
import time
import uuid
import re
import math
import pickle
from typing import List, Dict, Optional, Tuple

try:
    import numpy as np
    from sentence_transformers import SentenceTransformer
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    import pypdf
    import docx
    import easyocr
except ImportError as e:
    print(f"[WARN] RAG dependencies not installed yet: {e}")

class SimpleVectorStore:
    """
    一个简单的基于 Numpy 的向量检索存储
    替代 ChromaDB 以解决 SQLite 版本兼容性问题
    """
    def __init__(self, storage_path: str):
        self.storage_path = storage_path
        self.vectors_path = os.path.join(storage_path, "vectors.npy")
        self.meta_path = os.path.join(storage_path, "metadata.pkl")
        
        self.vectors = None # numpy array
        self.metadata = []  # list of dicts (aligned with vectors)
        
        self._load()

    def _load(self):
        if os.path.exists(self.vectors_path) and os.path.exists(self.meta_path):
            try:
                self.vectors = np.load(self.vectors_path)
                with open(self.meta_path, 'rb') as f:
                    self.metadata = pickle.load(f)
                print(f"[VectorStore] Loaded {len(self.metadata)} vectors.")
            except Exception as e:
                print(f"[VectorStore] Load failed: {e}")
                self.vectors = None
                self.metadata = []
        else:
            self.vectors = None
            self.metadata = []

    def _save(self):
        if not os.path.exists(self.storage_path):
            os.makedirs(self.storage_path, exist_ok=True)
            
        if self.vectors is not None:
            np.save(self.vectors_path, self.vectors)
            with open(self.meta_path, 'wb') as f:
                pickle.dump(self.metadata, f)

    def add(self, embeddings: List[List[float]], metadatas: List[Dict]):
        """添加向量"""
        new_vecs = np.array(embeddings, dtype=np.float32)
        
        if self.vectors is None:
            self.vectors = new_vecs
        else:
            self.vectors = np.vstack([self.vectors, new_vecs])
            
        self.metadata.extend(metadatas)
        self._save()

    def search(self, query_embedding: List[float], filter_fn=None, top_k: int = 3) -> List[Dict]:
        """
        搜索最相似的向量
        filter_fn: function(metadata) -> bool
        """
        if self.vectors is None or len(self.vectors) == 0:
            return []
            
        # Normalize vectors for cosine similarity
        # Cosine Similarity = (A . B) / (||A|| * ||B||)
        # Assuming query_embedding is not normalized
        
        q_vec = np.array(query_embedding, dtype=np.float32)
        q_norm = np.linalg.norm(q_vec)
        if q_norm == 0:
            return []
            
        # Calculate cosine similarity with all vectors
        # vectors shape: (N, D), q_vec shape: (D,)
        
        # Calculate norms for all storage vectors (cache this if slow)
        d_norms = np.linalg.norm(self.vectors, axis=1)
        
        # Dot product
        dot_products = np.dot(self.vectors, q_vec)
        
        # Similarity
        # Handle zero division
        d_norms[d_norms == 0] = 1e-10
        similarities = dot_products / (d_norms * q_norm)
        
        # Sort indices by similarity descending
        sorted_indices = np.argsort(similarities)[::-1]
        
        results = []
        count = 0
        
        for idx in sorted_indices:
            meta = self.metadata[idx]
            
            # Apply filter
            if filter_fn and not filter_fn(meta):
                continue
                
            results.append({
                "score": float(similarities[idx]),
                "metadata": meta,
                "id": meta.get("chunk_id") # Assuming chunk_id is in metadata
            })
            
            count += 1
            if count >= top_k:
                break
                
        return results

    def delete(self, filter_fn):
        """删除符合条件的向量"""
        if self.vectors is None:
            return
            
        indices_to_keep = []
        new_metadata = []
        
        for i, meta in enumerate(self.metadata):
            if not filter_fn(meta):
                indices_to_keep.append(i)
                new_metadata.append(meta)
                
        if len(indices_to_keep) < len(self.metadata):
            self.vectors = self.vectors[indices_to_keep]
            self.metadata = new_metadata
            self._save()

class KnowledgeBaseManager:
    def __init__(self, db_path: str = None, vector_db_path: str = None):
        # 1. SQL DB
        if db_path:
            from ai_expert.database import AIExpertDatabase
            self.sql_db = AIExpertDatabase(db_path)
        else:
            from ai_expert.database import AIExpertDatabase
            self.sql_db = AIExpertDatabase()
            
        # 2. Vector DB (Custom)
        if not vector_db_path:
            vector_db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'vector_store')
        
        self.vector_store = SimpleVectorStore(vector_db_path)

        # 3. Text & OCR
        try:
            # Use multilingual model for Chinese support
            self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
            print("[RAG] Embedding model loaded: paraphrase-multilingual-MiniLM-L12-v2")
        except Exception as e:
            print(f"[RAG] Model loading failed: {e}")
            self.model = None

        self.ocr_reader = None

    def _get_ocr_reader(self):
        if not self.ocr_reader:
            print("[RAG] Initializing EasyOCR...")
            self.ocr_reader = easyocr.Reader(['ch_sim', 'en'])
        return self.ocr_reader

    # ... (Add document logic same as before, but using self.vector_store.add)
    
    def add_document(self, file_path: str, bound_prompt_id: int = None, description: str = "") -> bool:
        """添加文档到知识库"""
        if not self.model: 
            return False

        if not os.path.exists(file_path): 
            return False

        # ... (Extract text - same as before) ...
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower().replace('.', '')
        text_content = ""
        try:
            if file_ext == 'pdf':
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
            elif file_ext in ['docx', 'doc']:
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                for table in doc.tables:
                    text_content += self._table_to_markdown(table) + "\n"
            elif file_ext in ['jpg', 'jpeg', 'png', 'bmp']:
                reader = self._get_ocr_reader()
                result = reader.readtext(file_path, detail=0)
                text_content = "\n".join(result)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
        except:
            return False

        if not text_content.strip(): return False

        # SQL Save
        try:
            conn = self.sql_db.get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO files (file_name, file_path, file_type, file_size, bound_prompt_id, description)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (file_name, file_path, file_ext, len(text_content), bound_prompt_id, description))
            file_id = cursor.lastrowid
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"[RAG] SQL Error: {e}")
            return False

        # Chunking
        chunks = self._chunk_text(text_content)
        
        # Embed and Save
        try:
            embeddings_list = self.model.encode(chunks)
        except:
            return False

        conn = self.sql_db.get_connection()
        cursor = conn.cursor()
        
        metadatas = []
        
        for i, chunk in enumerate(chunks):
            # Save to SQL chunks
            cursor.execute("""
                INSERT INTO chunks (file_id, chunk_index, content, token_count)
                VALUES (?, ?, ?, ?)
            """, (file_id, i, chunk, len(chunk)))
            
            meta = {
                "file_id": file_id,
                "bound_prompt_id": bound_prompt_id if bound_prompt_id is not None else 0,
                "chunk_index": i,
                "source": file_name,
                "content_preview": chunk[:50] # For debugging if needed
            }
            metadatas.append(meta)
            
        conn.commit()
        conn.close()
        
        # Save to Vector Store
        self.vector_store.add(embeddings_list.tolist(), metadatas)
        return True

    def search(self, query: str, bound_prompt_id: int = None, top_k: int = 3, threshold: float = 0.4) -> List[Dict]:
        """检索 (Threshold is Similarity threshold here, meaning min score)"""
        if not self.model: return []

        query_embedding = self.model.encode([query])[0].tolist()
        
        target_pid = bound_prompt_id if bound_prompt_id is not None else 0
        
        def filter_fn(meta):
            pid = meta.get("bound_prompt_id", 0)
            return pid == 0 or pid == target_pid

        results = self.vector_store.search(query_embedding, filter_fn, top_k)
        
        structured_results = []
        for res in results:
            # res: {score, metadata}
            # Threshold: if score < threshold (e.g. 0.4 similarity), skip
            # Note: Previously we used Distance. Here used Cosine Similarity.
            # Similarity is -1 to 1. 0.4 is reasonable.
            if res['score'] < threshold:
                continue
                
            # Retrieve full content from SQL or helper. 
            # Wait, I didn't save full content in metadata to save RAM, but I saved it in SQL.
            # But in this simple implementation, let's just fetch from SQL based on file_id/chunk_index 
            # OR better: fetch from SQL chunks table.
            
            # Efficient way:
            file_id = res['metadata']['file_id']
            chunk_index = res['metadata']['chunk_index']
            source = res['metadata']['source']
            
            content = self._get_chunk_content(file_id, chunk_index)
            
            structured_results.append({
                "content": content,
                "source": source,
                "score": res['score']
            })
            
        return structured_results

    def _get_chunk_content(self, file_id, chunk_index):
        conn = self.sql_db.get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT content FROM chunks WHERE file_id=? AND chunk_index=?", (file_id, chunk_index))
        row = cursor.fetchone()
        conn.close()
        return row['content'] if row else ""

    def _table_to_markdown(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append(f"| {' | '.join(cells)} |")
        if not rows: return ""
        header_len = len(table.rows[0].cells)
        separator = f"| {' | '.join(['---'] * header_len)} |"
        rows.insert(1, separator)
        return "\n" + "\n".join(rows) + "\n"

    def _chunk_text(self, text: str) -> List[str]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=100,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""]
        )
        return splitter.split_text(text)

    def delete_file(self, file_id: int) -> bool:
        # 1. SQL check
        conn = self.sql_db.get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM files WHERE id = ?", (file_id,))
        if not cursor.fetchone():
            conn.close()
            return False
            
        # 2. Vector delete
        def filter_fn(meta):
            return meta.get("file_id") == file_id
            
        self.vector_store.delete(filter_fn)
        
        # 3. SQL delete
        cursor.execute("DELETE FROM chunks WHERE file_id = ?", (file_id,))
        cursor.execute("DELETE FROM files WHERE id = ?", (file_id,))
        conn.commit()
        conn.close()
        return True

    def get_file_list(self, bound_prompt_id: int = None) -> List:
        conn = self.sql_db.get_connection()
        cursor = conn.cursor()
        if bound_prompt_id:
             cursor.execute("""
                SELECT * FROM files 
                WHERE bound_prompt_id = ? OR bound_prompt_id IS NULL OR bound_prompt_id = 0
                ORDER BY upload_time DESC
             """, (bound_prompt_id,))
        else:
             cursor.execute("SELECT * FROM files ORDER BY upload_time DESC")
        rows = cursor.fetchall()
        conn.close()
        return rows
